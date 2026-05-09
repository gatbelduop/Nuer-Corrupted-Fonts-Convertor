"""
Nuer Font Corruption Converter – Streamlit App
Converts corrupted Nuer fallback fonts to corrected Unicode Nuer.
Learning: user corrections saved and reused.
Numbers inside words (except at line start) can be optionally converted.
"""

# ======================================================================
# DEVELOPER CREDIT – Edit this section as you wish
# ======================================================================
# Developer Name:   Gatbel Duop Chol
# Website:          https://gatbelduop.github.io/Info/
# Contact:          gatbelduopchol@gmail.com
# Date:             2026
# ======================================================================
# To show credit in the app, change the next line to: SHOW_CREDIT_IN_UI = True
SHOW_CREDIT_IN_UI = True
# ======================================================================

import streamlit as st
import io
import zipfile
import threading
import time
import json
import re
from pathlib import Path
from typing import List, Tuple, Dict, Optional

# ----------------------------------------------------------------------
# Optional imports with graceful fallback
# ----------------------------------------------------------------------
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from odf.opendocument import load as load_odt
    from odf.text import P
    HAS_ODT = True
except ImportError:
    HAS_ODT = False

try:
    from striprtf.striprtf import rtf_to_text
    HAS_RTF = True
except ImportError:
    HAS_RTF = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from langdetect import detect
    HAS_LANGDETECT = True
except ImportError:
    HAS_LANGDETECT = False
    detect = None

# ----------------------------------------------------------------------
# MAPPING (derived from songs 1-38)
# ----------------------------------------------------------------------
CHAR_MAP = {
    # Lowercase
    ']': 'ɛ',     '[': 'ɔ',     '{': 'ŋ',     '}': 'ɣ',     '|': 'c',
    '@': 'a̠',     '=': 'a̱',     '>': 'ɛ',     ',': 'ɔ',     ';': 'i̠',
    '`': '’',     '~': 'ɣ',     '^': 'Ö',     '_': ' ',     '+': 'ö',
    '*': 'ä',     '&': 'ë',     '%': 'ɛ̈',     '#': 'ɔ̈',     '$': 'ä',
    '!': 'ï',     '\\': '/',    '/': '/',
    
    # Additional lowercase corrections (digits removed – handled separately)
    '`': 'ä',     
    ']': 'ɔ̱',     's': 'ɔ',
    'f': 'ɣ',     'x': 'ŋ',     'v': 'ɛ',
    
    # Uppercase mappings
    '~': 'Ä',     '!': 'A̱',     '@': 'Ë',     '#': 'E̱',     '%': 'I̱',
    '^': 'Ö',     '&': 'O̱',     ')': 'Ɛ̈',     '}': 'Ɔ̱',     'S': 'Ɔ',
    'F': 'Ɣ',     'X': 'Ŋ',     'V': 'Ɛ'
}

# Build translation table, excluding digits (they are handled separately)
TRANS_TABLE = str.maketrans({k: v for k, v in CHAR_MAP.items() if not k.isdigit()})

# Phrase mappings (some contain digits like 'j2' – these are handled first)
PHRASE_MAP = {
    "Ku,th": "Kuɔth", "Ku,]": "Kuɔɛ", "Ku=ar": "Kuäär",
    "G,,y": "Gɔɔy", "Ku``r": "Kuäär", "Nh=ok": "Nhök",
    "Bu,m": "Buɔm", "th=n": "thi̠n", "Th=n": "Thi̠n",
    "c],": "cɛ", "t=]thl,ac": "tɛ̈thlɔac", "T=]thl,ac": "Tɛ̈thlɔac",
    "l=ny": "lɔny", "L=ny": "Lɔny", "j2": "jɛ", "J2": "Jɛ",
    "n2": "nɛ", "N2": "Nɛ", "m2": "mɛ", "M2": "Mɛ",
    "v2": "ɛ", "V2": "Ɛɛ", "@i": "i̠", "@o": "o̠", "@u": "u̠",
    "@a": "a̠", "@e": "e̠", "@n": "i̠n", "ci=": "ci̠", "c=": "cä",
    "k=": "k̠", "p=": "p̠", "t=": "t̠", "r=": "r̠", "l=": "l̠",
    "Ki,th": "Kiɔth", "liak]": "liakɛ", "Th@in": "Thi̠n", "Ku[[r": "Kuäär",
    "Gssy": "Gɔɔy", "Nh]k": "Nhök", "vsaan2": "ŋɔaani̠",
    "wec muqqn": "wec muɔ̱n", "t\\thlsaac": "tɛ̈thlɔaac", "Yecu": "Yëcu",
    "Kritho": "Kri̠tho", "Alfluya": "Alɛluya", "Amen": "A-mɛn",
    "baa Ci=otdu": "baa Ciöötdu", "pu,ny": "puɔny", "Ku=ar>": "Kuäärä",
    "~=o=o": "ɣöö"
}

# ----------------------------------------------------------------------
# Digit mapping for inside-word conversion (separate simple mapping)
# Default values – can be customised by user in the sidebar.
# ----------------------------------------------------------------------
DEFAULT_DIGIT_MAP = {
    '0': 'ɔ',
    '1': 'ɛ',
    '2': 'ë',
    '3': 'e',
    '4': 'i',
    '5': 'ï',
    '6': 'ö',
    '7': 'o',
    '8': 'u',
    '9': 'a'
}

# Initialise session state for digit mapping (persists across reruns)
if 'digit_map' not in st.session_state:
    st.session_state.digit_map = DEFAULT_DIGIT_MAP.copy()
if 'digit_trans_table' not in st.session_state:
    st.session_state.digit_trans_table = str.maketrans(st.session_state.digit_map)

# English word list (common)
ENGLISH_WORDS = set([
    'the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have', 'i',
    'it', 'for', 'not', 'on', 'with', 'he', 'as', 'you', 'do', 'at',
    'this', 'but', 'his', 'by', 'from', 'they', 'we', 'say', 'her', 'she',
    'or', 'an', 'will', 'my', 'one', 'all', 'would', 'there', 'their',
    'what', 'so', 'up', 'out', 'if', 'about', 'who', 'get', 'which', 'go',
    'me', 'when', 'make', 'can', 'like', 'time', 'no', 'just', 'him', 'know'
])

# ----------------------------------------------------------------------
# User corrections (learning) with fast regex merging
# ----------------------------------------------------------------------
CORRECTIONS_FILE = Path(__file__).parent / "user_corrections.json"

def load_user_corrections() -> Dict[str, str]:
    if CORRECTIONS_FILE.exists():
        with open(CORRECTIONS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_user_correction(original: str, corrected: str):
    corrections = load_user_corrections()
    corrections[original] = corrected
    with open(CORRECTIONS_FILE, 'w', encoding='utf-8') as f:
        json.dump(corrections, f, indent=2, ensure_ascii=False)

def clear_all_corrections():
    if CORRECTIONS_FILE.exists():
        CORRECTIONS_FILE.unlink()
    st.success("All user corrections cleared. Refresh the app to reload changes.")
    time.sleep(1)
    st.rerun()

def _build_combined_regex(corrections: Dict[str, str]) -> re.Pattern:
    """Combine phrase map and user corrections into one regex (longest first)."""
    all_items = list(PHRASE_MAP.items()) + list(corrections.items())
    all_items.sort(key=lambda x: len(x[0]), reverse=True)
    pattern = '|'.join(re.escape(orig) for orig, _ in all_items)
    return re.compile(pattern, re.IGNORECASE)

# ----------------------------------------------------------------------
# Digit‑inside‑word conversion (with preservation at line start)
# ----------------------------------------------------------------------
def convert_digits_inside_words(text: str, preserve_line_start_digits: bool = True) -> str:
    """
    Convert digits that appear inside words.
    - Digits at the very beginning of a line are preserved if preserve_line_start_digits=True.
    - Words that consist only of digits are treated as a whole: if they are at line start -> preserved,
      otherwise each digit is converted individually.
    - Uses the digit mapping stored in st.session_state.
    """
    lines = text.splitlines(keepends=True)
    result_lines = []
    for line in lines:
        if preserve_line_start_digits:
            match = re.match(r'^([0-9]+)', line)
            if match:
                leading_digits = match.group(1)
                rest = line[len(leading_digits):]
                rest_converted = _convert_digits_in_text(rest)
                result_lines.append(leading_digits + rest_converted)
                continue
        result_lines.append(_convert_digits_in_text(line))
    return ''.join(result_lines)

def _convert_digits_in_text(text: str) -> str:
    """Replace all digits using the current digit translation table."""
    # Use session state mapping (updated by UI)
    trans = st.session_state.get('digit_trans_table', str.maketrans(DEFAULT_DIGIT_MAP))
    return text.translate(trans)

# ----------------------------------------------------------------------
# Fast English detection with caching
# ----------------------------------------------------------------------
_english_cache: Dict[str, bool] = {}

def is_english_word(word: str, use_langdetect: bool = False) -> bool:
    w = word.strip('.,!?;:()[]{}"\'')
    if not w or len(w) < 2:
        return False
    if w.lower() in ENGLISH_WORDS:
        return True
    if w.isalpha() and w.isascii():
        return True
    if use_langdetect and HAS_LANGDETECT:
        if w not in _english_cache:
            try:
                result = detect(w) == 'en'
            except:
                result = False
            _english_cache[w] = result
        return _english_cache[w]
    return False

# ----------------------------------------------------------------------
# Core conversion (with optional digit‑inside‑word conversion)
# ----------------------------------------------------------------------
def convert_text(text: str, preserve_english: bool = False, convert_digits: bool = False) -> str:
    """
    Convert corrupted Nuer text to proper Unicode.
    If preserve_english is True, attempt to keep English words unchanged.
    If convert_digits is True, apply digit‑inside‑word conversion (line-start digits preserved).
    """
    corrections = load_user_corrections()
    combined_regex = _build_combined_regex(corrections)

    def replacer(match):
        matched = match.group(0)
        for orig, corr in list(corrections.items()) + list(PHRASE_MAP.items()):
            if matched.lower() == orig.lower():
                if matched.isupper():
                    return corr.upper()
                elif matched[0].isupper() and matched[1:].islower():
                    return corr.capitalize()
                else:
                    return corr
        return matched

    text = combined_regex.sub(replacer, text)

    if preserve_english:
        tokens = re.split(r'(\s+|[.,!?;:()])', text)
        converted_tokens = []
        for token in tokens:
            if is_english_word(token, use_langdetect=False):
                converted_tokens.append(token)
            else:
                converted_tokens.append(token.translate(TRANS_TABLE))
        text = ''.join(converted_tokens)
    else:
        text = text.translate(TRANS_TABLE)

    if '@' in text:
        if any(c in 'aeiou' for c in text):
            text = text.replace('@', 'a̠')
        else:
            text = text.replace('@', 'i̠')

    if convert_digits:
        text = convert_digits_inside_words(text, preserve_line_start_digits=True)

    return text

# ----------------------------------------------------------------------
# File readers and writers (unchanged)
# ----------------------------------------------------------------------
def read_txt(data: bytes) -> str:
    return data.decode('utf-8', errors='replace')

def read_docx(data: bytes) -> str:
    if not HAS_DOCX:
        return "Error: python-docx not installed."
    doc = Document(io.BytesIO(data))
    return '\n'.join([para.text for para in doc.paragraphs])

def read_pdf(data: bytes) -> str:
    if not HAS_PDF:
        return "Error: PyPDF2 not installed."
    reader = PyPDF2.PdfReader(io.BytesIO(data))
    return '\n'.join([page.extract_text() for page in reader.pages])

def read_odt(data: bytes) -> str:
    if not HAS_ODT:
        return "Error: odfpy not installed."
    odt = load_odt(io.BytesIO(data))
    paras = [p.getAttribute('text') for p in odt.getElementsByType(P) if p.getAttribute('text')]
    return '\n'.join(paras)

def read_rtf(data: bytes) -> str:
    if not HAS_RTF:
        return "Error: striprtf not installed."
    return rtf_to_text(data.decode('utf-8', errors='replace'))

def write_txt(text: str) -> bytes:
    return text.encode('utf-8')

def write_docx(text: str) -> bytes:
    if not HAS_DOCX:
        return b"Error: python-docx not installed."
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def write_pdf(text: str) -> bytes:
    if not HAS_REPORTLAB:
        return b"Error: reportlab not installed."
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    y = height - 40
    for line in text.split('\n'):
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(40, y, line[:100])
        y -= 14
    c.save()
    return bio.getvalue()

# ----------------------------------------------------------------------
# Batch conversion job
# ----------------------------------------------------------------------
class ConversionJob:
    def __init__(self, files: List[Tuple[str, bytes]], output_format: str, preserve_english: bool, convert_digits: bool):
        self.files = files
        self.output_format = output_format
        self.preserve_english = preserve_english
        self.convert_digits = convert_digits
        self.results = []
        self.progress = 0
        self.is_done = False
        self.error = None

    def run(self):
        try:
            total = len(self.files)
            for i, (name, data) in enumerate(self.files):
                ext = Path(name).suffix.lower()
                if ext == '.txt':
                    text = read_txt(data)
                elif ext == '.docx':
                    text = read_docx(data)
                elif ext == '.pdf':
                    text = read_pdf(data)
                elif ext == '.odt':
                    text = read_odt(data)
                elif ext == '.rtf':
                    text = read_rtf(data)
                else:
                    text = data.decode('utf-8', errors='replace')
                converted = convert_text(text, self.preserve_english, self.convert_digits)
                out_name = Path(name).stem + self.output_format
                if self.output_format == '.txt':
                    out_data = write_txt(converted)
                elif self.output_format == '.docx':
                    out_data = write_docx(converted)
                elif self.output_format == '.pdf':
                    out_data = write_pdf(converted)
                else:
                    out_data = converted.encode('utf-8')
                self.results.append((out_name, out_data))
                self.progress = (i + 1) / total
            self.is_done = True
        except Exception as e:
            self.error = str(e)
            self.is_done = True

# ----------------------------------------------------------------------
# Streamlit UI
# ----------------------------------------------------------------------
st.set_page_config(page_title="Nuer Font Converter", layout="wide")
st.title("📝 Nuer Font Corruption Converter")

st.sidebar.warning(
    "⚠️ **Important limitation**\n\n"
    "This converter is **not perfect**. Many corrupted words may remain "
    "or be incorrectly converted.\n\n"
    "**Please help improve it!** Use the **✏️ Help improve the converter** "
    "section below the converted text to add your own corrections.\n\n"
    "Your contributions will be saved and reused for everyone."
)

if SHOW_CREDIT_IN_UI:
    st.sidebar.markdown("---")
    st.sidebar.info("Developed by Gatbel Duop Chol – gatbelduopchol@gmail.com - https://gatbelduop.github.io/Info/")

with st.sidebar:
    st.markdown("---")
    st.subheader("⚙️ Options")
    output_format = st.selectbox("Output format", [".txt", ".docx", ".pdf"], index=0)
    preserve_english = st.checkbox(
        "🔤 Preserve English words (slower for large files)",
        value=False,
        help="Disable for maximum speed when converting pure Nuer text."
    )
    convert_digits = st.checkbox(
        "🔢 Convert digits inside words (experimental)",
        value=False,
        help="Convert digits that appear inside Nuer words (e.g., 'th2n' → 'thën'). Digits at the beginning of a line are preserved."
    )
    st.markdown("---")
    st.markdown("### 📚 How it works")
    st.markdown(
        "This tool repairs text corrupted by missing Nuer font support. "
        "It uses character mappings and learned corrections.\n\n"
        "**Digit conversion rule:** Digits inside words are replaced using a separate mapping. "
        "Digits at the very start of a line remain unchanged. Words that are only digits are preserved only if they start a line."
    )
    
    # Editable digit mapping (stored in session state)
    with st.expander("✏️ Edit digit‑inside‑word mapping"):
        st.markdown("Current mapping (digit → character):")
        new_map = {}
        for d in '0123456789':
            current = st.session_state.digit_map.get(d, DEFAULT_DIGIT_MAP.get(d, ''))
            new_val = st.text_input(f"Digit {d}", value=current, key=f"digit_{d}")
            new_map[d] = new_val if new_val else current
        if st.button("Update digit mapping"):
            st.session_state.digit_map.update(new_map)
            st.session_state.digit_trans_table = str.maketrans(st.session_state.digit_map)
            st.success("Digit mapping updated!")
            st.rerun()
    
    st.markdown("---")
    st.subheader("📋 User Corrections")
    corrections = load_user_corrections()
    if corrections:
        st.json(corrections)
    else:
        st.info("No user corrections saved yet.")
    if st.button("🗑️ Clear all corrections", use_container_width=True):
        clear_all_corrections()

# Main area: Input methods
tab1, tab2 = st.tabs(["✏️ Text input", "📂 Batch files"])

with tab1:
    input_text = st.text_area("Paste corrupted text here:", height=300)
    if st.button("Convert", key="convert_text"):
        if input_text.strip():
            with st.spinner("Converting..."):
                result = convert_text(input_text, preserve_english, convert_digits)
                st.session_state['last_result'] = result
        else:
            st.warning("Please enter some text.")

with tab2:
    uploaded_files = st.file_uploader(
        "Upload files (TXT, DOCX, PDF, ODT, RTF)", 
        accept_multiple_files=True,
        type=['txt', 'docx', 'pdf', 'odt', 'rtf']
    )
    batch_mode = st.checkbox("Batch mode (download as ZIP)", value=True)
    if uploaded_files and st.button("Convert files", key="convert_files"):
        valid_files = []
        max_size_mb = 50
        for f in uploaded_files:
            if f.size <= max_size_mb * 1024 * 1024:
                valid_files.append((f.name, f.getvalue()))
            else:
                st.warning(f"{f.name} exceeds {max_size_mb} MB – skipped.")
        if valid_files:
            job = ConversionJob(valid_files, output_format, preserve_english, convert_digits)
            thread = threading.Thread(target=job.run)
            thread.start()
            prog_bar = st.progress(0)
            while not job.is_done:
                prog_bar.progress(job.progress)
                time.sleep(0.05)
            prog_bar.empty()
            if job.error:
                st.error(job.error)
            else:
                if batch_mode:
                    zip_bio = io.BytesIO()
                    with zipfile.ZipFile(zip_bio, 'w') as zf:
                        for name, dat in job.results:
                            zf.writestr(name, dat)
                    st.download_button("📦 Download all as ZIP", zip_bio.getvalue(),
                                       file_name="converted_batch.zip", mime="application/zip")
                else:
                    for name, dat in job.results:
                        st.download_button(f"📥 {name}", dat, file_name=name)
        else:
            st.warning("No valid files provided.")

# Display result and learning interface
if 'last_result' in st.session_state:
    st.subheader("✅ Converted text")
    st.text_area("Result", st.session_state['last_result'], height=400)
    col1, col2 = st.columns(2)
    with col1:
        st.download_button("💾 Download as TXT", st.session_state['last_result'].encode('utf-8'),
                           file_name="converted.txt", mime="text/plain")
    with col2:
        st.markdown("**✏️ Help improve the converter**")
        orig_word = st.text_input("Original corrupted word/phrase (exactly as above):")
        correct_word = st.text_input("Corrected version (proper Nuer):")
        if st.button("Submit correction") and orig_word and correct_word:
            save_user_correction(orig_word, correct_word)
            st.success(f"Saved: {orig_word} → {correct_word}. Restart app to use updated mapping.")
            st.info("The regex will be rebuilt automatically on next conversion.")
